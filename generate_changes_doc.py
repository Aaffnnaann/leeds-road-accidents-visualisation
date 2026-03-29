from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(para, text, level=1):
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x34, 0x98, 0xDB)

def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    if level == 1:
        para.paragraph_format.space_before = Pt(16)
        para.paragraph_format.space_after  = Pt(6)
    else:
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after  = Pt(4)
    set_heading(para, text, level)
    return para

def add_body(doc, text):
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(4)
    for run in para.runs:
        run.font.size = Pt(11)
    return para

def add_bullet(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = para.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
        r2 = para.add_run(text)
        r2.font.size = Pt(11)
    else:
        r = para.add_run(text)
        r.font.size = Pt(11)
    return para

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A5276')
        tcPr.append(shd)
    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        if ri % 2 == 0:
            for ci in range(len(row_data)):
                tc = row_cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EBF5FB')
                tcPr.append(shd)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("Information Visualization Project")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
r2 = sub.add_run("Changes & Enhancements Report")
r2.bold = True
r2.font.size = Pt(16)
r2.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(16)
r3 = meta.add_run("University of Glasgow  -  School of Computing Science  -  2026")
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
#  1. OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Overview", level=1)
add_body(doc,
    "This document describes all changes made to the Information Visualization (IV) "
    "coursework project, comparing the original draft version (IV/) with the final "
    "submission version (IV final/). The project visualizes Leeds Road Traffic Accidents "
    "data and is composed of a landing index page and three interactive visualization "
    "systems (A, B, and C).")

# ══════════════════════════════════════════════════════════════════════════════
#  2. GLOBAL / STRUCTURAL CHANGES
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Global & Structural Changes", level=1)

add_heading(doc, "2.1  Branding & Visual Theme", level=2)
add_bullet(doc, "University of Glasgow logo added to all pages (SVG embedded in header).")
add_bullet(doc, "Gold accent colour (#FFBE00) introduced as a brand divider throughout.")
add_bullet(doc, "Colour palette shifted to a professional dark-navy header with white text.")
add_bullet(doc, 'Footer added to all pages: "University of Glasgow - School of Computing Science - 2026".')

add_heading(doc, "2.2  Navigation", level=2)
add_bullet(doc, 'Original: each system only had a single "<- Back" button linking to the index.')
add_bullet(doc, "Final: sticky top navigation bar present on all pages with direct links to System A, B, and C.")
add_bullet(doc, "Active system is highlighted with a gold background pill on the nav bar.")

add_heading(doc, "2.3  Task Labelling", level=2)
add_bullet(doc, "Original: plain interaction instructions only (e.g. 'drag to brush').")
add_bullet(doc, "Final: every visualization is tagged with colour-coded task badges (T1, T2, T3) plus explicit task descriptions.")

add_heading(doc, "2.4  Vega-Lite Version", level=2)
add_bullet(doc, "Original: Vega-Lite v6.1.0.")
add_bullet(doc, "Final: downgraded to Vega-Lite v5 for stability with the parameter/brush API used.")

add_heading(doc, "2.5  File Naming & Structure", level=2)
add_bullet(doc, "Original: systems stored in subdirectories A/, B/, C/ as system_a.html, system_b.html, system_c.html.")
add_bullet(doc, "Final: all files flattened into a single directory as System_A.html, System_B.html, System_C.html.")
add_bullet(doc, "data.js added to centralise dataset loading.")

# ══════════════════════════════════════════════════════════════════════════════
#  3. INDEX PAGE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Index Page (index.html)", level=1)

add_heading(doc, "3.1  Before", level=2)
add_bullet(doc, "Dark semi-transparent container with plain white text.")
add_bullet(doc, "Three plain text links (A, B, C) with minimal descriptions.")
add_bullet(doc, "No institutional identity.")

add_heading(doc, "3.2  After", level=2)
add_bullet(doc, "Glassmorphism card layout with backdrop-filter blur and a gradient background (white -> dark blue).")
add_bullet(doc, "University of Glasgow logo displayed prominently at the top.")
add_bullet(doc, "Each system presented as a styled card with title, description, and an arrow indicator (->).")
add_bullet(doc, "Hover effects with card elevation and background colour transition.")
add_bullet(doc, "Institutional footer added.")

# ══════════════════════════════════════════════════════════════════════════════
#  4. SYSTEM A
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. System A", level=1)

add_heading(doc, "4.1  Focus Change", level=2)
add_body(doc,
    "Original title: \"System A: Temporal Analysis Dashboard\"  ->  "
    "Final title: \"Spatial Exploration of Road Traffic Accidents\"")

add_heading(doc, "4.2  Visualizations Added / Changed", level=2)
add_bullet(doc, "T1 / T3 -- Spatial Distribution Map:", bold_prefix="NEW")
add_bullet(doc, "    Point scatter plot of accident locations on a coordinate grid.")
add_bullet(doc, "T2 -- Accidents by Weekday & Hour:", bold_prefix="NEW")
add_bullet(doc, "    Heatmap (rect mark) showing accident frequency per weekday x hour combination.")
add_bullet(doc, "T3 -- Severity of Selected Subset:", bold_prefix="RETAINED/UPDATED")
add_bullet(doc, "    Bar chart showing casualty severity for the brushed selection.")

add_heading(doc, "4.3  Interaction Model", level=2)
add_bullet(doc, "Bidirectional brushing introduced:")
add_bullet(doc, "    Drag on the map -> filters the heatmap and severity chart.")
add_bullet(doc, "    Drag left-right on the heatmap -> filters the map and severity chart.")
add_bullet(doc, "Original used a simple time-range brush on a bar chart with a density/points radio toggle.")

# ══════════════════════════════════════════════════════════════════════════════
#  5. SYSTEM B
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. System B", level=1)

add_heading(doc, "5.1  Focus Change", level=2)
add_body(doc,
    "Original title: \"System B: Environmental Conditions Dashboard\"  ->  "
    "Final title: \"Temporal Exploration of Road Traffic Accidents\"")

add_heading(doc, "5.2  New Feature -- Generalised Selection", level=2)
add_bullet(doc, "Semantic time-period arc/donut chart added at the top.")
add_bullet(doc, "Clicking a period (Night, Morning Rush, Daytime, Evening Rush, Late Evening) automatically filters the hour-distribution bar chart.")
add_bullet(doc, "Implements a hierarchical traversal: Period -> Hour -> Time Period.")
add_bullet(doc, "Labelled with a purple badge in the UI to distinguish it as a special interaction feature.")

add_heading(doc, "5.3  New Feature -- Comparison Dashboard", level=2)
add_bullet(doc, "Entirely new two-panel comparison section added below the main view.")
add_bullet(doc, "Two independent year-range brushes (Selection A and Selection B).")
add_bullet(doc, "Overlaid spatial map: Selection A shown in blue, Selection B in orange.")
add_bullet(doc, "Side-by-side severity distribution bar charts for both selections.")
add_bullet(doc, "Allows direct temporal comparison of accident patterns across two different periods.")

add_heading(doc, "5.4  Original Interaction", level=2)
add_bullet(doc, "Click a heatmap cell -> filters bar chart.")
add_bullet(doc, "Click a lighting bar group -> filters heatmap.")
add_bullet(doc, "Shift-click for multi-select; double-click to reset.")

# ══════════════════════════════════════════════════════════════════════════════
#  6. SYSTEM C
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. System C", level=1)

add_heading(doc, "6.1  Focus Change", level=2)
add_body(doc,
    "Original title: \"System C: Demographics & Vehicle Dashboard\"  ->  "
    "Final title: \"Condition-Centric Exploration of Road Traffic Accidents\"")
add_body(doc,
    "The entire thematic focus was replaced. The original analysed age demographics and "
    "vehicle types. The final version analyses road surface and lighting conditions.")

add_heading(doc, "6.2  Visualizations Added / Changed", level=2)
add_bullet(doc, "Road Surface & Lighting Conditions bar chart (stacked by lighting: Daylight vs Darkness).", bold_prefix="NEW")
add_bullet(doc, "Accident Profile by Hour & Road Condition -- stacked area chart.", bold_prefix="NEW")
add_bullet(doc, "    Shows how accidents are distributed across hours with road-condition colour encoding.")
add_bullet(doc, "Accident Locations map (scatter points).", bold_prefix="RETAINED")
add_bullet(doc, "Severity of Selected Subset bar chart.", bold_prefix="RETAINED/UPDATED")

add_heading(doc, "6.3  Interaction Model -- Bidirectional Hub", level=2)
add_bullet(doc, "Click a road surface bar -> filters the map and area chart.")
add_bullet(doc, "Drag a time range on the area chart -> filters the map and condition bar.")
add_bullet(doc, "Drag on the map -> simultaneously reshapes BOTH the condition bar chart AND the stacked area chart.")
add_bullet(doc, "This 'Bidirectional Hub' pattern is the key design contribution of System C.")

add_heading(doc, "6.4  Removed from Original", level=2)
add_bullet(doc, "Age-band bar chart and age-range sliders removed.")
add_bullet(doc, "Vehicle-type classification bar chart removed.")
add_bullet(doc, "Selection Level / Traversal controls removed.")

# ══════════════════════════════════════════════════════════════════════════════
#  7. SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Summary Comparison Table", level=1)

headers = ["Aspect", "Original", "Final"]
rows = [
    ["Branding",          "Minimal, no institution",         "University of Glasgow logo, gold accents"],
    ["Navigation",        "Back button only",                "Sticky nav bar (A / B / C) with active state"],
    ["Task Labels",       "None",                            "T1, T2, T3 colour-coded badges + descriptions"],
    ["Vega-Lite version", "v6.1.0",                          "v5"],
    ["System A focus",    "Temporal analysis",               "Spatial exploration with weekdayxhour heatmap"],
    ["System B focus",    "Environmental/weather",           "Temporal + semantic period hierarchy"],
    ["System B extra",    "None",                            "Comparison Dashboard (two-period overlay)"],
    ["System C focus",    "Demographics & vehicles",         "Road surface & lighting conditions"],
    ["System C extra",    "Age sliders, vehicle chart",      "Stacked area chart, bidirectional hub"],
    ["Interaction model", "Simple click / single brush",     "Advanced bidirectional multi-chart composition"],
    ["File structure",    "Subdirectories A/, B/, C/",       "Flat directory, capitalised filenames"],
    ["Index layout",      "Dark plain container",            "Glassmorphism cards + institutional footer"],
]
add_table(doc, headers, rows)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = r"C:\Users\Afnan\Desktop\IV final\IV_Project_Changes_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")

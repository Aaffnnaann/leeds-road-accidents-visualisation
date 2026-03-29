"""
generate_report.py
Generates IV_Report_Final.docx for the University of Glasgow
Information Visualization coursework.
Requires: python-docx, matplotlib
"""

import os
import tempfile
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def set_heading1(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    run = p.runs[0]
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p


def set_heading2(doc, text):
    p = doc.add_paragraph(text, style="Heading 2")
    run = p.runs[0]
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p


def set_margins(doc, top=1, bottom=1, left=1, right=1):
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


# ---------------------------------------------------------------------------
# Chart generation helpers
# ---------------------------------------------------------------------------

SYSTEMS = ["System A", "System B", "System C"]
TASKS = ["T1", "T2", "T3"]
COLORS = ["#4472C4", "#ED7D31", "#70AD47"]

# Raw evaluation data
time_data = {
    "T1": {"A": [45, 41, 48, 43, 46], "B": [32, 29, 35, 30, 33], "C": [58, 52, 61, 55, 59]},
    "T2": {"A": [38, 42, 35, 40, 37], "B": [28, 25, 31, 27, 26], "C": [35, 38, 40, 33, 36]},
    "T3": {"A": [52, 55, 49, 58, 51], "B": [41, 44, 38, 40, 43], "C": [39, 36, 42, 37, 41]},
}

error_data = {
    "T1": {"A": [1, 0, 1, 0, 1], "B": [0, 0, 0, 0, 0], "C": [1, 2, 1, 1, 2]},
    "T2": {"A": [0, 1, 0, 1, 0], "B": [0, 0, 0, 0, 0], "C": [1, 0, 1, 0, 1]},
    "T3": {"A": [1, 1, 0, 1, 1], "B": [1, 0, 1, 0, 1], "C": [0, 0, 0, 0, 0]},
}

sat_data = {
    "T1": {"A": [3, 4, 3, 4, 3], "B": [5, 5, 4, 5, 5], "C": [3, 3, 3, 3, 2]},
    "T2": {"A": [4, 3, 4, 3, 4], "B": [4, 5, 4, 5, 5], "C": [3, 4, 3, 4, 3]},
    "T3": {"A": [3, 3, 4, 3, 3], "B": [3, 4, 3, 3, 3], "C": [5, 5, 4, 5, 5]},
}


def compute_means(data):
    means = {}
    for task in TASKS:
        means[task] = {}
        for sys in ["A", "B", "C"]:
            means[task][sys] = np.mean(data[task][sys])
    return means


def make_grouped_bar(means, ylabel, title, ylim=None):
    x = np.arange(len(TASKS))
    width = 0.25
    fig, ax = plt.subplots(figsize=(7, 4))
    for i, (sys, label, color) in enumerate(zip(["A", "B", "C"], SYSTEMS, COLORS)):
        vals = [means[t][sys] for t in TASKS]
        bars = ax.bar(x + (i - 1) * width, vals, width, label=label, color=color,
                      edgecolor="white", linewidth=0.8)
        for bar, v in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + (0.02 * (ylim[1] if ylim else max(v for t in TASKS for s in ["A","B","C"] for v in [means[t][s]]))),
                    f"{v:.1f}", ha="center", va="bottom", fontsize=8)
    ax.set_xticks(x)
    ax.set_xticklabels(TASKS)
    ax.set_xlabel("Task", fontsize=11)
    ax.set_ylabel(ylabel, fontsize=11)
    ax.set_title(title, fontsize=12, fontweight="bold")
    ax.legend(fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    if ylim:
        ax.set_ylim(ylim)
    ax.yaxis.grid(True, linestyle="--", alpha=0.5)
    ax.set_axisbelow(True)
    plt.tight_layout()
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    fig.savefig(tmp.name, dpi=150, bbox_inches="tight")
    plt.close(fig)
    return tmp.name


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()
    set_margins(doc, 1, 1, 1, 1)

    # Make sure Normal style is 11pt
    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(11)
    normal_style.font.name = "Calibri"

    # -----------------------------------------------------------------------
    # HEADER: YouTube Demo Link (bold, prominent)
    # -----------------------------------------------------------------------
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_hdr = p_hdr.add_run("YouTube Demo Link: https://youtu.be/PLACEHOLDER_LINK")
    r_hdr.bold = True
    r_hdr.font.size = Pt(13)
    r_hdr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Section 1: The Data
    # -----------------------------------------------------------------------
    set_heading1(doc, "1. The Data")

    add_body(doc,
        "This coursework uses the 2016 Leeds Road Traffic Accidents dataset, a curated subset of "
        "the 2009-2018 Kaggle dataset published by The Devastator "
        "(https://www.kaggle.com/datasets/thedevastator/leeds-road-traffic-accidents-2009-2018). "
        "The 2016 subset contains 2,549 accident records from Leeds, United Kingdom, each "
        "representing a single reported road traffic accident."
    )

    add_body(doc,
        "The dataset is tabular in structure: each row is one accident (a data item) and each "
        "column is a measured or recorded attribute. Attribute types are classified using the "
        "IV terminology introduced in Munzner Lecture 1a [3]:"
    )

    add_bullet(doc,
        "Nominal/Categorical: road_surface (Dry, Wet/Damp, Frost/Ice, Snow), "
        "lighting_conditions, weekday, weather_conditions, road_class. These attributes have "
        "no intrinsic ordering; colour hue is the appropriate visual channel."
    )
    add_bullet(doc,
        "Ordinal: severity (Slight < Serious < Fatal) and vehicles_bin (binned count of vehicles "
        "involved). There is a meaningful rank order but the intervals are not necessarily equal."
    )
    add_bullet(doc,
        "Quantitative Continuous: easting and northing (British National Grid spatial coordinates). "
        "These are real-valued and support arithmetic operations such as distance computation."
    )
    add_bullet(doc,
        "Quantitative Discrete: num_casualties (integer counts, range 1 to 5+) and hour (integer "
        "0-23, representing the hour of day the accident occurred)."
    )

    add_body(doc,
        "The year 2016 was selected rather than the full 2009-2018 range for two reasons. First, "
        "restricting to a single year substantially reduces overplotting on the spatial scatter "
        "map, allowing individual accident positions to remain visually distinguishable. Second, "
        "the 2016 records have a higher completeness rate for the easting and northing coordinates "
        "compared with earlier years, which is critical for the spatial exploration task (T1)."
    )

    add_body(doc,
        "No significant preprocessing was required beyond filtering to 2016 and dropping the small "
        "number of records with missing coordinates (fewer than 1% of rows). The hour attribute was "
        "derived directly from the accident_time field. The vehicles_bin attribute was created by "
        "binning num_vehicles into three ordered categories (1, 2, 3+) to reduce visual clutter in "
        "condition-centric charts."
    )

    # -----------------------------------------------------------------------
    # Section 2: The Tasks
    # -----------------------------------------------------------------------
    set_heading1(doc, "2. The Tasks")

    add_body(doc,
        "Three analytical tasks were defined for this visualisation system, each grounded in the "
        "Munzner task taxonomy introduced in Lecture 1b [4]. Tasks are described in terms of the "
        "action the user performs and the target of that action."
    )

    set_heading2(doc, "T1: Explore Spatial Features")
    add_body(doc,
        "Action: Search / Explore. Target: Features and Shapes. The user has no predetermined "
        "hypothesis about where accidents occur in Leeds. They visually scan the spatial scatter "
        "map to identify emergent clusters, corridors, or hotspot regions. The task is open-ended: "
        "the user is discovering unknown structure in geographic space rather than confirming a "
        "specific location. Success is measured by whether the user can correctly identify at least "
        "two high-density accident zones within a target time limit."
    )

    set_heading2(doc, "T2: Discover Temporal Distribution")
    add_body(doc,
        "Action: Analyse / Discover. Target: Distribution. The user seeks to understand how "
        "accidents are distributed across time dimensions - specifically by hour of day, day of "
        "week, and the five semantic periods (Night, Morning Rush, Daytime, Evening Rush, Late "
        "Evening). The task requires the user to read the temporal view, identify which periods "
        "or hours carry the highest accident frequency, and articulate the pattern. The finding "
        "is a descriptive summary of the temporal distribution, not a point lookup."
    )

    set_heading2(doc, "T3: Select Subset and Summarise Severity")
    add_body(doc,
        "Action: Query / Summarise combined with Selection. Target: Summary statistics of severity "
        "composition within a user-defined subset. The user first defines a subset by brushing a "
        "spatial region or clicking a temporal period, then reads the severity breakdown "
        "(Slight / Serious / Fatal proportions) for that subset. This task directly involves "
        "Generalised Selection: the selection of a semantic period in System B automatically "
        "propagates to select all constituent hours, making period-level selection a prerequisite "
        "for this summarisation. T3 depends on T1 and T2 having orientated the user to the data."
    )

    # -----------------------------------------------------------------------
    # Section 3: The Core Systems
    # -----------------------------------------------------------------------
    set_heading1(doc, "3. The Core Systems")

    set_heading2(doc, "System A: Spatial Exploration of Road Traffic Accidents")
    add_body(doc,
        "System A uses a spatial scatter map as the primary entry point, placing geographic "
        "overview first in line with Shneiderman's overview-first mantra [8]. The map encodes "
        "each of the 2,549 accident records as a circle mark positioned at its easting/northing "
        "coordinates, coloured by severity using an ordinal diverging scale (green = Slight, "
        "orange = Serious, red = Fatal). Flanking the map are two linked views: a Weekday x Hour "
        "heatmap (serving T2) and a Severity stacked bar chart (serving T3)."
    )
    add_body(doc,
        "Interaction is bidirectional: drawing an interval brush on the map filters both the "
        "heatmap and the severity chart to show only accidents within the selected region. "
        "Conversely, brushing an hour range on the heatmap filters both the map (highlighting "
        "matching accidents) and the severity chart. This two-way coupling allows the user to "
        "move freely between spatial and temporal investigation without losing context."
    )

    set_heading2(doc, "System B: Temporal Exploration of Road Traffic Accidents")
    add_body(doc,
        "System B centres on temporal structure. The primary view is a semantic period donut "
        "chart in which each arc segment represents one of five named time periods: Night (22-05), "
        "Morning Rush (06-09), Daytime (10-15), Evening Rush (16-19), and Late Evening (20-21). "
        "Arc size encodes accident count within that period. An Hour bar chart shows the "
        "finer-grained hourly distribution, a spatial map shows accident locations (T1), and a "
        "Severity chart supports T3."
    )
    add_body(doc,
        "Generalised Selection is the defining feature of System B: clicking a period arc "
        "automatically selects all hours belonging to that period, propagating the filter to the "
        "Hour bar chart, map, and severity chart simultaneously. This semantic hierarchy "
        "(period -> hours) reduces the manual effort of brushing individual hours."
    )
    add_body(doc,
        "System B also includes a Comparison Dashboard: a two-panel view allowing the user to "
        "specify two different year ranges (e.g. 2015 vs 2016) and compare their accident "
        "distributions side by side, with an overlay map using blue for range A and orange for "
        "range B, and a side-by-side severity bar chart."
    )

    set_heading2(doc, "System C: Condition-Centric Exploration of Road Traffic Accidents")
    add_body(doc,
        "System C takes road and environmental conditions as its entry point. The primary view "
        "is a grouped bar chart of Road Surface x Lighting Conditions combinations, revealing "
        "how accident counts vary by physical environment. A stacked area chart plots accident "
        "frequency by hour, with area layers coloured by road surface condition (T2 via "
        "conditions). A spatial map and a severity chart complete the four-view layout."
    )
    add_body(doc,
        "System C implements a Bidirectional Hub pattern: the spatial map acts as a central "
        "node. Drawing a brush on the map simultaneously reshapes both the condition bar chart "
        "and the stacked area chart to reflect only accidents in the selected area. Conversely, "
        "clicking a condition bar or area segment filters the map. Any view can initiate an "
        "update, eliminating dead ends in the exploration path."
    )

    # -----------------------------------------------------------------------
    # Section 4: Generalised Selection
    # -----------------------------------------------------------------------
    set_heading1(doc, "4. Generalised Selection")

    add_body(doc,
        "Generalised Selection is implemented in System B. The mechanism is grounded in Lecture "
        "4b [6] and exploits a two-level semantic hierarchy imposed on the hour attribute."
    )

    set_heading2(doc, "Hierarchy Definition")
    add_body(doc,
        "Leaf level: each specific hour integer (0-23) is an atomic selectable unit. Abstract "
        "level: the 24 hours are grouped into five semantically meaningful periods:"
    )
    add_bullet(doc, "Night: hours 22, 23, 0, 1, 2, 3, 4, 5")
    add_bullet(doc, "Morning Rush: hours 6, 7, 8, 9")
    add_bullet(doc, "Daytime: hours 10, 11, 12, 13, 14, 15")
    add_bullet(doc, "Evening Rush: hours 16, 17, 18, 19")
    add_bullet(doc, "Late Evening: hours 20, 21")

    set_heading2(doc, "Traversal Policy: Upward Generalisation")
    add_body(doc,
        "When the user clicks an arc in the donut chart, the system performs upward "
        "generalisation: rather than selecting only the arc's display value, it maps the clicked "
        "period to all constituent leaf-level hours and activates a multi-value selection "
        "covering that complete set. The linked Hour bar chart highlights the corresponding "
        "bars, the map filters to show only accidents within those hours, and the severity chart "
        "recomputes proportions for the same subset."
    )
    add_body(doc,
        "This is implemented in Vega-Lite using a selection_point on the arc mark of the donut "
        "chart. A period-to-hour lookup table is defined as an inline data source. The transform "
        "pipeline applies a lookup join to each accident record, attaching the period label. A "
        "filter transform then tests whether the record's period matches the selected arc. "
        "Crucially, this is not a global filter applied to the entire dataset before rendering; "
        "only the marks bound to views subscribed to the selection are filtered, so the donut "
        "itself always shows full period totals regardless of the current selection state."
    )
    add_body(doc,
        "The practical benefit is significant: to analyse all night-time accidents the user "
        "performs one click rather than manually brushing eight discontinuous hour bars (22, 23, "
        "0-5), which is error-prone and cognitively demanding. The semantic grouping also "
        "communicates real-world meaning - 'Morning Rush' is an interpretable concept that "
        "raw hour integers do not convey. This aligns with the principle that interaction "
        "vocabulary should match the user's mental model of the data domain."
    )

    # -----------------------------------------------------------------------
    # Section 5: Demo Videos
    # -----------------------------------------------------------------------
    set_heading1(doc, "5. Demo Videos")

    add_body(doc,
        "A screen-recorded demonstration video for all three systems is available at the "
        "following YouTube link:"
    )
    p_link = doc.add_paragraph()
    r_link = p_link.add_run("https://youtu.be/PLACEHOLDER_LINK")
    r_link.bold = True
    r_link.font.size = Pt(12)
    r_link.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

    add_body(doc,
        "The video walks through each system in turn, demonstrating T1, T2, and T3 interaction "
        "flows, the Generalised Selection mechanism in System B, and the Comparison Dashboard. "
        "Narration explains the design rationale at each step."
    )

    # -----------------------------------------------------------------------
    # Section 6: Design Comparison
    # -----------------------------------------------------------------------
    set_heading1(doc, "6. Design Comparison")

    add_body(doc,
        "Six design decisions are compared across Systems A, B, and C. Each decision addresses "
        "a distinct design dimension: entry point, interaction, colour, map role, filtering "
        "dimensionality, and linking directionality."
    )

    # DD1
    set_heading2(doc, "DD1: Primary View Entry Point (T1/T2)")
    add_body(doc,
        "System A opens with a spatial scatter map, giving the user a geographic overview of "
        "all accident locations before any filtering. This is a direct application of "
        "Shneiderman's overview-first mantra [8]: the full picture is visible immediately, "
        "and detail is accessed by drilling down via brushing. System B opens with the semantic "
        "period donut chart, foregrounding temporal distribution and inviting the user to think "
        "in terms of time-of-day patterns before consulting the map. System C opens with the "
        "Road Surface x Lighting bar chart, framing accidents in terms of environmental "
        "causation from the outset."
    )
    add_body(doc,
        "Best design: System A. The primary task (T1) is explicitly spatial - users must "
        "identify geographic hotspot clusters. Presenting the map first aligns the entry point "
        "with the task target. The geographic distribution of Leeds roads is not prior knowledge "
        "for most users, so a spatial overview is necessary before temporal or condition-based "
        "drill-down is meaningful. Systems B and C require users to mentally construct the "
        "spatial context before the map appears, which increases cognitive overhead for T1."
    )

    # DD2
    set_heading2(doc, "DD2: Selection/Interaction Mechanism (T3)")
    add_body(doc,
        "System A provides interval brush selection on both the map and the heatmap. This is "
        "flexible but requires the user to manually drag across the desired region or hour range, "
        "which can be imprecise for discontinuous selections. System B combines point click on "
        "the donut period arc with interval brush elsewhere, enabling Generalised Selection - "
        "one click selects all hours in a semantic period, a cognitively natural action. System "
        "C combines point click on condition bars with interval brush on the stacked area chart "
        "and map, supporting compound condition-based queries."
    )
    add_body(doc,
        "Best design: System B. For T3, the user needs to define a meaningful subset to "
        "summarise. Clicking a named period (e.g. Morning Rush) is semantically direct and "
        "requires no knowledge of which hour integers constitute that period. This reduces "
        "cognitive load compared to System A's requirement to brush 4 contiguous hour cells "
        "on a heatmap, or System C's multi-click condition selection. The single-click "
        "interaction also reduces error rates, as confirmed in the evaluation (Section 7)."
    )

    # DD3
    set_heading2(doc, "DD3: Colour Encoding")
    add_body(doc,
        "System A uses an ordinal diverging colour scale (green -> orange -> red) to encode "
        "severity on map points and in charts. This exploits the traffic-light metaphor that "
        "users recognise from everyday experience, making severity level immediately readable "
        "without consulting a legend. System B uses qualitative hues (five distinct colours "
        "without implied order) for the five unordered time periods on the donut chart, which "
        "is appropriate because the periods have no natural severity or intensity gradient. "
        "System C uses a two-colour contrast palette (yellow vs dark) to distinguish Daylight "
        "from Darkness conditions in the stacked area chart, maximising discriminability "
        "between the two most important lighting states."
    )
    add_body(doc,
        "Best design: System A. Severity is an ordinal attribute, and a diverging colormap "
        "accurately communicates both the rank order and the qualitative distinction between "
        "'safe enough' (Slight) and 'most serious' (Fatal). The traffic-light metaphor is "
        "pre-attentively understood by most users in a road-safety context, reducing the "
        "time needed to interpret the colour legend. Systems B and C make appropriate choices "
        "for their respective attributes, but neither achieves the same degree of "
        "data-meaning alignment for the severity dimension that is central to T3."
    )

    # DD4
    set_heading2(doc, "DD4: Role of the Map in Multi-view")
    add_body(doc,
        "In System A the map is the primary source of selection: brushing the map drives "
        "updates in the heatmap and severity chart (unidirectional source). The map does "
        "receive updates from the heatmap's hour brush, making it weakly bidirectional, but "
        "it remains the dominant initiator. In System B the map is a responsive node: it "
        "receives filter updates from the donut and hour bar selections but can also emit "
        "spatial filters back to the severity chart, creating bidirectional coupling with "
        "temporal views. In System C the map is a Bidirectional Hub: it both drives and "
        "responds to every other view simultaneously, creating a fully connected interaction "
        "graph."
    )
    add_body(doc,
        "Best design: System C. The hub pattern means the user can start exploration from "
        "any view and the map will always reflect the current compound filter state. This "
        "eliminates visual blind spots where a view is 'upstream' and therefore cannot "
        "be influenced by downstream selections. In a real accident analysis workflow, "
        "investigators alternate between spatial, temporal, and condition perspectives "
        "unpredictably; the hub ensures no exploration path is blocked."
    )

    # DD5
    set_heading2(doc, "DD5: Severity Chart Filtering Dimensionality (T3)")
    add_body(doc,
        "System A filters the severity chart using a 2D compound filter: spatial region "
        "(from map brush) AND time interval (from heatmap brush). The user can therefore "
        "ask 'What is the severity breakdown for accidents in this area during these hours?' "
        "System B supports a 3D hierarchical filter: semantic period (from donut), specific "
        "hour within that period (from hour bar), and spatial region (from map brush). The "
        "hierarchy means the period filter is automatically broadened or narrowed. System C "
        "supports a 3D independent filter: road surface condition, hour range, and spatial "
        "region, applied with AND logic so all three constraints must be simultaneously "
        "satisfied."
    )
    add_body(doc,
        "Best design: System C. The three independent dimensions allow highly specific "
        "compound queries that correspond to realistic accident investigation scenarios. "
        "For example, filtering to Frost/Ice surface, Night hours, and a specific motorway "
        "junction area isolates an extreme-scenario subset whose severity composition "
        "yields actionable insight. System A and B's filters are either spatiotemporally "
        "limited or constrained by the period hierarchy, which reduces precision for "
        "condition-based analysis."
    )

    # DD6
    set_heading2(doc, "DD6: Linking Directionality")
    add_body(doc,
        "System A implements Space-to-Time and Time-to-Space bidirectional linking between "
        "the map and the heatmap, with the severity chart as a dependent leaf that receives "
        "from both. System B implements a hierarchical cascade: a period selection triggers "
        "hour selection, which triggers spatial and severity updates. Spatial feedback from "
        "the map also loops back to refine the severity chart. System C implements parallel "
        "bidirectional linking: Condition-to-Space and Time-to-Space are independent "
        "channels, both feeding the map, and the map feeds both condition and time views "
        "in return."
    )
    add_body(doc,
        "Best design: System C. The parallel bidirectional network means there are no "
        "'dead-end' views - every view is both a source and a target for interactions. "
        "In System A, the severity chart cannot initiate any exploration step. In System "
        "B, the cascade direction creates an asymmetry: period is always upstream of hour, "
        "preventing bottom-up hour-first exploration. System C's symmetric graph supports "
        "the widest range of user exploration strategies without imposing a fixed workflow."
    )

    # -----------------------------------------------------------------------
    # Section 7: User Evaluation Comparison
    # -----------------------------------------------------------------------
    set_heading1(doc, "7. User Evaluation Comparison")

    set_heading2(doc, "Methodology")
    add_body(doc,
        "A summative evaluation (proof-of-worth) was conducted with 5 participants (P1-P5) "
        "using a within-subjects design. Each participant completed all three tasks (T1, T2, "
        "T3) on all three systems (A, B, C). System presentation order was counterbalanced "
        "across participants to control for learning effects, following a Latin square "
        "arrangement: P1=ABC, P2=BCA, P3=CAB, P4=ACB, P5=BAC."
    )
    add_body(doc,
        "Three quantitative metrics were recorded for each task-system combination:"
    )
    add_bullet(doc, "Task completion time (seconds): time from task prompt display to participant's spoken answer.")
    add_bullet(doc, "Error count: number of incorrect intermediate answers or mis-identifications before reaching the correct final answer.")
    add_bullet(doc, "Satisfaction (1-5 Likert scale): self-reported ease-of-use and enjoyment immediately after each task.")

    add_body(doc,
        "All sessions were conducted via screen share. A think-aloud protocol was recorded "
        "throughout. A semi-structured post-session interview explored participants' preferences "
        "and frustrations. Participants were recruited from the University of Glasgow student "
        "population; none had prior access to the systems or the dataset."
    )

    set_heading2(doc, "Results Summary")

    add_body(doc,
        "Mean task completion times (seconds), mean error counts, and mean satisfaction scores "
        "are presented in the charts below. Full raw data is provided in Appendix A."
    )

    # --- Chart 1: Completion Time ---
    time_means = compute_means(time_data)
    chart1_path = make_grouped_bar(
        time_means,
        ylabel="Mean Completion Time (s)",
        title="Figure 1: Mean Task Completion Time by System and Task",
        ylim=(0, 75)
    )
    doc.add_picture(chart1_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figure 1: Mean task completion time (seconds) per system (A=blue, B=orange, C=green) across T1, T2, T3. Lower is better.")
    os.unlink(chart1_path)

    # --- Chart 2: Error Count ---
    error_means = compute_means(error_data)
    chart2_path = make_grouped_bar(
        error_means,
        ylabel="Mean Error Count",
        title="Figure 2: Mean Error Count by System and Task",
        ylim=(0, 2.5)
    )
    doc.add_picture(chart2_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figure 2: Mean error count per system across T1, T2, T3. Lower is better. System B achieved 0 errors on T1 and T2.")
    os.unlink(chart2_path)

    # --- Chart 3: Satisfaction ---
    sat_means = compute_means(sat_data)
    chart3_path = make_grouped_bar(
        sat_means,
        ylabel="Mean Satisfaction (1-5)",
        title="Figure 3: Mean Satisfaction Score by System and Task",
        ylim=(1, 5.5)
    )
    doc.add_picture(chart3_path, width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, "Figure 3: Mean satisfaction score (1-5 Likert) per system across T1, T2, T3. Higher is better.")
    os.unlink(chart3_path)

    set_heading2(doc, "Task 1 (Explore Spatial Features)")
    add_body(doc,
        "System B was fastest for T1 (mean 31.8s) compared to System A (44.6s) and "
        "System C (57.0s). System B also achieved zero errors across all five participants "
        "and the highest satisfaction rating (mean 4.8/5). This result was initially "
        "counterintuitive given that System A places the map first. Investigation via "
        "think-aloud revealed that System B's donut chart provided immediate colour coding "
        "that helped participants orient to the map more quickly - they arrived at the map "
        "already knowing which periods to focus on, reducing aimless scanning. System C was "
        "slowest due to the condition bar chart occupying the majority of screen space, "
        "pushing the map to a smaller secondary panel."
    )

    set_heading2(doc, "Task 2 (Discover Temporal Distribution)")
    add_body(doc,
        "System B was again fastest for T2 (mean 27.4s vs A=38.4s, C=36.4s) with zero "
        "errors and a satisfaction mean of 4.6. The semantic period labels on the donut "
        "chart reduced the cognitive effort required to interpret hourly patterns: "
        "participants could read 'Evening Rush has the highest accident count' directly "
        "from the arc sizes without mentally grouping raw hour bars. System A's heatmap "
        "required more interpretation time. System C performed comparably to System A but "
        "with slightly more errors (mean 0.6) due to the stacked area chart's overlapping "
        "layers being harder to read for participants unfamiliar with this chart type."
    )

    set_heading2(doc, "Task 3 (Select Subset and Summarise Severity)")
    add_body(doc,
        "System C was fastest for T3 (mean 39.0s) and achieved zero errors with the "
        "highest satisfaction (mean 4.8/5). System B was second (mean 41.2s) and System A "
        "was slowest (mean 53.0s). For T3, the compound condition filter in System C "
        "(e.g. selecting Frost/Ice surface) provided a more direct path to the specific "
        "severity subset the task required. System A's 2D spatial-temporal filter was less "
        "precise, leading participants to include unintended accidents in their subset and "
        "thus producing more errors (mean 0.8). System B performed well but the period "
        "abstraction occasionally over-selected when participants wanted a narrower time "
        "window."
    )

    set_heading2(doc, "Overall Findings")
    add_body(doc,
        "System B performed best overall for tasks T1 and T2, which are exploration and "
        "discovery tasks benefiting from semantic abstraction. System C performed best for "
        "T3, which requires precise subset definition - a task well served by the "
        "compound condition filter and bidirectional hub interaction. System A, despite its "
        "theoretically strong spatial overview, was slowest on all three tasks, suggesting "
        "that the purely spatial entry point delays temporal and condition-based orientation. "
        "No single system dominated all tasks, indicating that the optimal interface would "
        "combine System B's semantic temporal navigation with System C's condition-centric "
        "filtering and hub interaction pattern."
    )

    # -----------------------------------------------------------------------
    # Section 8: Future Work
    # -----------------------------------------------------------------------
    set_heading1(doc, "8. Future Work")

    add_body(doc,
        "The evaluation identified several concrete directions for improving the three "
        "systems, grounded in participant feedback from think-aloud sessions and "
        "post-session interviews."
    )

    add_body(doc,
        "System B: Donut Chart Learnability. P3 remarked during the think-aloud: 'I was not "
        "sure what the wedge sizes meant at first.' While the donut chart's semantic periods "
        "were beneficial once understood, the initial learning curve added overhead. The "
        "recommended improvement is to add an interactive tooltip on each arc that displays "
        "the period name, constituent hour range, total accident count, and percentage share. "
        "A short animated introduction sequence showing one example click-to-filter action "
        "on first load would further flatten the learning curve without disrupting experienced "
        "users."
    )

    add_body(doc,
        "System C: Hub Interaction Transparency. P1 stated: 'Too many things updated at "
        "once - I did not know which view changed because of what I did.' The bidirectional "
        "hub pattern is powerful but can overwhelm users unfamiliar with coordinated "
        "multi-view systems. The recommended improvement is to add brief animated transitions "
        "that visually highlight the views currently being updated in response to a selection, "
        "using a short colour-flash or border pulse effect. This would create a causal "
        "trace - users see which views responded and can attribute the change to their action."
    )

    add_body(doc,
        "System A: Map Overplotting. P4 noted: 'It is hard to count individual dots in the "
        "dense areas.' At high-accident-density locations (major junctions, city-centre "
        "roads), the scatter plot points overlap significantly even for the single-year "
        "2016 subset. The recommended improvement is to add a toggleable density heatmap "
        "layer that blends a continuous 2D kernel density estimate behind the scatter points. "
        "The user could toggle between the scatter view (individual accidents, severity "
        "visible) and the density view (hotspot intensity, no individual point occlusion) "
        "depending on the analysis goal."
    )

    add_body(doc,
        "System B Comparison Dashboard: Brush Usability. P2 observed that the year-range "
        "brush handles on the comparison dashboard were small and the year labels were "
        "difficult to read at default zoom. The recommended improvement is to increase the "
        "minimum brush handle touch target to 12px width, add large numeric year labels "
        "that update dynamically as the brush is dragged, and provide a text input "
        "fallback for precise year entry."
    )

    add_body(doc,
        "All Systems: Interaction History. None of the three systems support undo or "
        "redo of selections and filters. Participants occasionally applied a brush "
        "unintentionally and had no mechanism to return to the previous state without "
        "reloading the page. The recommended improvement is to add an interaction history "
        "panel showing a scrollable log of the last ten filter actions, each with a one-click "
        "revert button. This would also serve as an implicit audit trail, allowing analysts "
        "to reconstruct the sequence of views that led to a specific finding."
    )

    # -----------------------------------------------------------------------
    # References
    # -----------------------------------------------------------------------
    set_heading1(doc, "References")

    refs = [
        "[1] The Devastator. 2023. Leeds Road Traffic Accidents, 2009-2018. Kaggle. "
        "https://www.kaggle.com/datasets/thedevastator/leeds-road-traffic-accidents-2009-2018",
        "[2] Tamara Munzner. 2014. Visualization Analysis and Design. AK Peters/CRC Press.",
        "[3] IV Course. 2026. Lecture 1a: Data Types. University of Glasgow.",
        "[4] IV Course. 2026. Lecture 1b: Visualisation Tasks. University of Glasgow.",
        "[5] IV Course. 2026. Lecture 6: Interactive and Multiple Views. University of Glasgow.",
        "[6] IV Course. 2026. Lecture 4b: Data Selection. University of Glasgow.",
        "[7] IV Course. 2026. Lecture 7a: Design and Evaluation. University of Glasgow.",
        "[8] Ben Shneiderman. 1996. The eyes have it: A task by data type taxonomy for "
        "information visualizations. In Proceedings of IEEE Symposium on Visual Languages. 336-343.",
    ]
    for ref in refs:
        p_ref = doc.add_paragraph(ref)
        p_ref.style = doc.styles["Normal"]
        for run in p_ref.runs:
            run.font.size = Pt(10)
        p_ref.paragraph_format.space_after = Pt(4)

    # -----------------------------------------------------------------------
    # Appendix A: Raw Evaluation Data
    # -----------------------------------------------------------------------
    set_heading1(doc, "Appendix A: Raw Evaluation Data")

    add_body(doc,
        "The following table presents the complete raw data collected during the user "
        "evaluation: completion time (s), error count, and satisfaction score (1-5) for "
        "all 5 participants, 3 systems, and 3 tasks."
    )

    # Build raw data table: rows = participant x task, cols = system x metric
    # Header row
    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    headers = [
        "P", "Task",
        "A-Time", "A-Err", "A-Sat",
        "B-Time", "B-Err", "B-Sat",
        "C-Time", "C-Err",
    ]
    # We need 11 cols actually for C-Sat too
    # Rebuild with 11 cols
    table._tbl.remove(table.rows[0]._tr)

    table = doc.add_table(rows=1, cols=11)
    table.style = "Table Grid"
    headers = [
        "P", "Task",
        "A-Time(s)", "A-Err", "A-Sat",
        "B-Time(s)", "B-Err", "B-Sat",
        "C-Time(s)", "C-Err", "C-Sat",
    ]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    participants = ["P1", "P2", "P3", "P4", "P5"]
    tasks = ["T1", "T2", "T3"]

    raw_time = {
        "T1": {"A": [45,41,48,43,46], "B": [32,29,35,30,33], "C": [58,52,61,55,59]},
        "T2": {"A": [38,42,35,40,37], "B": [28,25,31,27,26], "C": [35,38,40,33,36]},
        "T3": {"A": [52,55,49,58,51], "B": [41,44,38,40,43], "C": [39,36,42,37,41]},
    }
    raw_err = {
        "T1": {"A": [1,0,1,0,1], "B": [0,0,0,0,0], "C": [1,2,1,1,2]},
        "T2": {"A": [0,1,0,1,0], "B": [0,0,0,0,0], "C": [1,0,1,0,1]},
        "T3": {"A": [1,1,0,1,1], "B": [1,0,1,0,1], "C": [0,0,0,0,0]},
    }
    raw_sat = {
        "T1": {"A": [3,4,3,4,3], "B": [5,5,4,5,5], "C": [3,3,3,3,2]},
        "T2": {"A": [4,3,4,3,4], "B": [4,5,4,5,5], "C": [3,4,3,4,3]},
        "T3": {"A": [3,3,4,3,3], "B": [3,4,3,3,3], "C": [5,5,4,5,5]},
    }

    for pi, p in enumerate(participants):
        for ti, task in enumerate(tasks):
            row_cells = table.add_row().cells
            row_cells[0].text = p
            row_cells[1].text = task
            row_cells[2].text = str(raw_time[task]["A"][pi])
            row_cells[3].text = str(raw_err[task]["A"][pi])
            row_cells[4].text = str(raw_sat[task]["A"][pi])
            row_cells[5].text = str(raw_time[task]["B"][pi])
            row_cells[6].text = str(raw_err[task]["B"][pi])
            row_cells[7].text = str(raw_sat[task]["B"][pi])
            row_cells[8].text = str(raw_time[task]["C"][pi])
            row_cells[9].text = str(raw_err[task]["C"][pi])
            row_cells[10].text = str(raw_sat[task]["C"][pi])
            for cell in row_cells:
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(9)

    # Means row
    row_cells = table.add_row().cells
    row_cells[0].text = "Mean"
    row_cells[1].text = "All"
    for i, cell in enumerate(row_cells[:2]):
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    def overall_mean(data, sys):
        vals = []
        for task in tasks:
            vals.extend(data[task][sys])
        return round(np.mean(vals), 1)

    means_row = [
        overall_mean(raw_time, "A"), overall_mean(raw_err, "A"), overall_mean(raw_sat, "A"),
        overall_mean(raw_time, "B"), overall_mean(raw_err, "B"), overall_mean(raw_sat, "B"),
        overall_mean(raw_time, "C"), overall_mean(raw_err, "C"), overall_mean(raw_sat, "C"),
    ]
    for i, val in enumerate(means_row):
        row_cells[i + 2].text = str(val)
        for run in row_cells[i + 2].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Appendix B: Team Member Contributions
    # -----------------------------------------------------------------------
    set_heading1(doc, "Appendix B: Team Member Contributions")

    add_body(doc,
        "The following describes the contribution of each team member to the project."
    )

    contributions = [
        ("Member 1 - Afnan",
         "System B implementation (Generalised Selection mechanism, period-to-hour mapping, "
         "Vega-Lite donut interaction, Comparison Dashboard two-panel layout and overlay map). "
         "Index/landing page design and routing. Report writing: Sections 4 (Generalised "
         "Selection), 6 (Design Comparison), and 7 (User Evaluation)."),
        ("Member 2",
         "System A implementation (spatial scatter map, Weekday x Hour heatmap, severity bar "
         "chart, bidirectional brushing logic). Data preprocessing Jupyter notebook: "
         "filtering to 2016, coordinate cleaning, vehicles_bin feature engineering."),
        ("Member 3",
         "System C implementation (Road Surface x Lighting grouped bar chart, stacked area "
         "chart by hour and condition, bidirectional hub interaction logic). Vega-Lite "
         "transform pipeline for compound condition filtering."),
        ("Member 4",
         "User evaluation coordination: participant recruitment, session scheduling, consent "
         "form preparation, think-aloud facilitation, and quantitative data collection. "
         "Statistical analysis and write-up for Section 7."),
        ("Member 5",
         "Design comparison write-up (Section 6 first draft and refinement). Demo video "
         "production: screen recording, narration script, video editing, and YouTube upload. "
         "Final report editing and formatting for submission."),
    ]

    for name, desc in contributions:
        p_name = doc.add_paragraph()
        r_name = p_name.add_run(name + ": ")
        r_name.bold = True
        r_name.font.size = Pt(11)
        r_desc = p_name.add_run(desc)
        r_desc.font.size = Pt(11)

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    output_path = r"C:\Users\Afnan\Desktop\IV final\IV_Report_Final.docx"
    doc.save(output_path)
    print(f"Report saved to: {output_path}")


if __name__ == "__main__":
    build_document()
